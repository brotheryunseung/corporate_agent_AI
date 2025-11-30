from docx import Document
from docx.shared import Pt

def export_to_word(report_text: str, filename: str = "equity_research_report.docx"):
    """
    Convert the analyst report text into a properly formatted Word document (.docx).
    """

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Add text line-by-line
    for line in report_text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    # Save file
    doc.save(filename)
    return filename