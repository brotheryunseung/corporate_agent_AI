




import yfinance as yf
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#########################################
# Utility functions
#########################################

def find_line(df: pd.DataFrame, keywords):
    if df is None or df.empty:
        return None, None

    idx_candidates = [
        idx for idx in df.index
        if any(k.lower() in str(idx).lower() for k in keywords)
    ]
    if not idx_candidates:
        return None, None

    row = df.loc[idx_candidates[0]]
    latest = row.iloc[0] if len(row) > 0 else None
    prev = row.iloc[1] if len(row) > 1 else None
    return latest, prev


def pct(x):
    try:
        return f"{x:.2%}"
    except:
        return "N/A"


def num(x):
    try:
        return f"{x:,.0f}"
    except:
        return "N/A"


#########################################
# Corporate Analysis (Main Report)
#########################################

def corporate_full_analysis(ticker: str) -> str:
    tk = yf.Ticker(ticker)

    try:
        info = tk.get_info()
    except Exception:
        info = {}

    company_name = info.get("longName", ticker)
    sector = info.get("sector", "N/A")
    industry = info.get("industry", "N/A")
    market_cap = info.get("marketCap", None)
    beta = info.get("beta", None)

    fin = tk.financials
    bs = tk.balance_sheet
    cf = tk.cashflow

    if fin is None or fin.empty:
        return f"❌ Could not retrieve financial statements for {ticker} using yfinance."

    revenue, prev_revenue = find_line(fin, ["Total Revenue", "Revenue"])
    gross_profit, _ = find_line(fin, ["Gross Profit"])
    op_income, _ = find_line(fin, ["Operating Income"])
    net_income, _ = find_line(fin, ["Net Income"])
    eps, prev_eps = find_line(fin, ["Basic EPS", "EPS", "Diluted EPS"])

    total_assets, _ = find_line(bs, ["Total Assets"])
    total_liab, _ = find_line(bs, ["Total Liabilities"])
    equity, _ = find_line(bs, ["Total Equity", "Total Stockholder Equity"])

    op_cf, _ = find_line(cf, ["Net Cash Provided by Operating Activities"])
    capex, _ = find_line(cf, ["Capital Expenditures"])
    free_cf = op_cf + capex if op_cf is not None and capex is not None else None

    def safe_div(a, b):
        try:
            if a is None or b is None or b == 0:
                return None
            return a / b
        except:
            return None

    revenue_growth = safe_div(revenue, prev_revenue) - 1 if (revenue and prev_revenue) else None
    eps_growth = safe_div(eps, prev_eps) - 1 if (eps and prev_eps) else None

    gross_margin = safe_div(gross_profit, revenue)
    op_margin = safe_div(op_income, revenue)
    net_margin = safe_div(net_income, revenue)

    roe = safe_div(net_income, equity)
    roa = safe_div(net_income, total_assets)
    debt_to_equity = safe_div(total_liab, equity)

    pe = info.get("trailingPE", None)
    forward_pe = info.get("forwardPE", None)
    ps = info.get("priceToSalesTrailing12Months", None)
    pb = info.get("priceToBook", None)

    report = f"""
==============================
📊 Corporate Analyst Report: {ticker}
==============================

🏢 Company Overview
Name: {company_name}
Sector: {sector}
Industry: {industry}
Market Cap: {num(market_cap)}
Beta: {beta if beta is not None else "N/A"}

💵 Profitability
Gross Margin: {pct(gross_margin)}
Operating Margin: {pct(op_margin)}
Net Margin: {pct(net_margin)}

📈 Growth
Revenue (latest): {num(revenue)}
Revenue Growth (YoY): {pct(revenue_growth)}
EPS (latest): {eps if eps is not None else "N/A"}
EPS Growth (YoY): {pct(eps_growth)}

📊 Returns
ROE: {pct(roe)}
ROA: {pct(roa)}

🔐 Financial Health
Total Assets: {num(total_assets)}
Total Liabilities: {num(total_liab)}
Total Equity: {num(equity)}
Debt-to-Equity: {format(debt_to_equity, '.2f') if debt_to_equity is not None else "N/A"}

🏭 Cash Flow
Operating Cash Flow: {num(op_cf)}
Capital Expenditures: {num(capex)}
Free Cash Flow: {num(free_cf)}

💰 Valuation
Trailing P/E: {pe}
Forward P/E: {forward_pe}
Price/Sales (TTM): {ps}
Price/Book: {pb}

==============================
Note: Automatically generated corporate analysis for personal use.
==============================
"""

    return report


#########################################
# Rating Engine
#########################################

def compute_core_metrics(ticker: str) -> dict:
    tk = yf.Ticker(ticker)
    
    try:
        info = tk.get_info()
    except:
        info = {}

    pe = info.get("trailingPE", None)
    ps = info.get("priceToSalesTrailing12Months", None)

    fin = tk.financials
    bs = tk.balance_sheet

    if fin is None or fin.empty:
        return {"ok": False, "reason": "No income statement found"}

    revenue, prev_revenue = find_line(fin, ["Total Revenue"])
    net_income, prev_net_income = find_line(fin, ["Net Income"])
    eps, prev_eps = find_line(fin, ["Basic EPS", "EPS", "Diluted EPS"])

    total_assets, _ = find_line(bs, ["Total Assets"])
    total_liab, _ = find_line(bs, ["Total Liabilities"])
    equity, _ = find_line(bs, ["Total Equity"])

    def safe_div(a, b):
        try:
            if not a or not b:
                return None
            return a / b
        except:
            return None

    revenue_growth = safe_div(revenue, prev_revenue) - 1 if (revenue and prev_revenue) else None
    eps_growth = safe_div(eps, prev_eps) - 1 if (eps and prev_eps) else None

    roe = safe_div(net_income, equity)
    roa = safe_div(net_income, total_assets)
    de = safe_div(total_liab, equity)

    return {
        "ok": True,
        "revenue_growth": revenue_growth,
        "eps_growth": eps_growth,
        "roe": roe,
        "roa": roa,
        "de": de,
        "pe": pe,
        "ps": ps
    }


def score_from_metrics(m: dict) -> dict:
    if not m["ok"]:
        return {"rating": "N/A", "score": 0, "components": {}}

    score = 0
    comp = {}

    # Profitability
    if m["roe"] is not None:
        comp["profit"] = 25 if m["roe"] > 0.15 else 15
    else:
        comp["profit"] = 10

    # Growth
    comp["growth"] = 25 if m["revenue_growth"] and m["revenue_growth"] > 0.1 else 10

    # Risk
    comp["risk"] = 20 if m["de"] and m["de"] < 1.0 else 5

    # Valuation
    comp["valuation"] = 20 if m["pe"] and m["pe"] < 25 else 10

    total = sum(comp.values())

    if total >= 75:
        rating = "BUY"
    elif total >= 50:
        rating = "HOLD"
    else:
        rating = "SELL"

    return {
        "rating": rating,
        "score": total,
        "components": comp
    }


#########################################
# Word Report Generator
#########################################

def create_word_report(ticker: str, analysis_text: str, rating_data: dict, file_name=None):
    if file_name is None:
        file_name = f"{ticker}_report.docx"
    
    doc = Document()
    
    title = doc.add_heading(f"{ticker} – Corporate Analyst Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = rating_data
    doc.add_paragraph(f"Investment Rating: {r['rating']}  |  Score: {r['score']} / 100")

    doc.add_paragraph()
    for line in analysis_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(file_name)
    return file_name